# ==========================================
# 1. IMPORTACIONES Y CONFIGURACIÓN INICIAL
# ==========================================
import streamlit as st
import pandas as pd
import sqlite3
import socket
import ssl
import datetime
import hashlib
import requests
import io
import base64
import re
import os
import csv
from urllib.parse import urlparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import matplotlib.pyplot as plt
import psycopg2
from weasyprint import HTML

st.set_page_config(
    page_title="CyberAudits - Escáner Perimetral Enterprise",
    page_icon="🛡️",
    layout="wide"
)

st.markdown("""
    <style>
        .stApp { background-color: #f8fafc; color: #1e293b; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; }
        [data-testid="stSidebar"] { background-color: #f0f2f6 !important; border-right: 1px solid #e2e8f0; }
        [data-testid="stSidebar"] label, [data-testid="stSidebar"] .stMarkdown, [data-testid="stSidebar"] span, [data-testid="stSidebar"] p, [data-testid="stSidebar"] h2 { color: #1e293b !important; }
        [data-testid="stSidebar"] input, [data-testid="stSidebar"] div[data-baseweb="select"] > div { background-color: #ffffff !important; color: #1e293b !important; border-color: #cbd5e1 !important; }
        .enterprise-banner { background: linear-gradient(90deg, #1e3a8a, #3b82f6); padding: 12px 20px; border-radius: 8px; color: white; text-align: center; margin-bottom: 20px; font-weight: 500; }
        .training-card { background: #ffffff; border: 1px solid #cbd5e1; border-left: 4px solid #3b82f6; padding: 20px; border-radius: 6px; margin-bottom: 20px; line-height: 1.6; }
        .employee-portal-banner { background: linear-gradient(90deg, #0f172a, #1e3a8a); padding: 20px; border-radius: 8px; color: white; text-align: center; margin-bottom: 25px; }
        .ticket-card { background: white; border: 1px solid #cbd5e1; border-radius: 8px; padding: 16px; margin-bottom: 12px; box-shadow: 0 1px 3px rgba(0,0,0,0.05); }
        .sev-critical { border-left: 5px solid #dc2626; }
        .sev-medium { border-left: 5px solid #f59e0b; }
        .sev-low { border-left: 5px solid #3b82f6; }
    </style>
""", unsafe_allow_html=True)


# ==========================================
# 2. GESTIÓN DE BASE DE DATOS
# ==========================================
def get_db_connection():
    if "postgres" in st.secrets and "url" in st.secrets["postgres"]:
        conn = psycopg2.connect(st.secrets["postgres"]["url"])
        conn.autocommit = True
        return conn
    else:
        conn = sqlite3.connect("cyber_audits.db")
        return conn

def init_db():
    conn = get_db_connection()
    conn.autocommit = True
    c = conn.cursor()
    is_pg = "postgres" in st.secrets
    
    if is_pg:
        c.execute("""CREATE TABLE IF NOT EXISTS organizations (id SERIAL PRIMARY KEY, name TEXT UNIQUE NOT NULL, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)""")
        c.execute("""CREATE TABLE IF NOT EXISTS history (id SERIAL PRIMARY KEY, timestamp TEXT, hostname TEXT, ip TEXT, risk_score INTEGER, findings_count INTEGER, report_type TEXT, organization_id INTEGER)""")
        c.execute("ALTER TABLE history ADD COLUMN IF NOT EXISTS organization_id INTEGER;")
        c.execute("""CREATE TABLE IF NOT EXISTS employees (id SERIAL PRIMARY KEY, organization_id INTEGER, email TEXT NOT NULL, department TEXT, topic TEXT DEFAULT 'Módulo 1 — Phishing', status TEXT DEFAULT 'Pendiente', score INTEGER DEFAULT 0, last_completed TEXT, UNIQUE(email, topic))""")
        c.execute("""CREATE TABLE IF NOT EXISTS remediation_tasks (id SERIAL PRIMARY KEY, organization_id INTEGER, scan_id INTEGER, hostname TEXT, finding_vector TEXT, severity TEXT DEFAULT 'MEDIO', status TEXT DEFAULT 'Pendiente', notes TEXT)""")
        c.execute("ALTER TABLE remediation_tasks ADD COLUMN IF NOT EXISTS organization_id INTEGER;")
        c.execute("ALTER TABLE remediation_tasks ADD COLUMN IF NOT EXISTS scan_id INTEGER;")
        c.execute("ALTER TABLE remediation_tasks ADD COLUMN IF NOT EXISTS severity TEXT;")
        c.execute("""CREATE TABLE IF NOT EXISTS remediation_logs (id SERIAL PRIMARY KEY, task_id INTEGER, timestamp TEXT, status TEXT, notes TEXT)""")
    else:
        c.execute("""CREATE TABLE IF NOT EXISTS organizations (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT UNIQUE NOT NULL, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)""")
        c.execute("""CREATE TABLE IF NOT EXISTS history (id INTEGER PRIMARY KEY AUTOINCREMENT, timestamp TEXT, hostname TEXT, ip TEXT, risk_score INTEGER, findings_count INTEGER, report_type TEXT, organization_id INTEGER)""")
        try: c.execute("ALTER TABLE history ADD COLUMN organization_id INTEGER;")
        except Exception: pass
        c.execute("""CREATE TABLE IF NOT EXISTS employees (id INTEGER PRIMARY KEY AUTOINCREMENT, organization_id INTEGER, email TEXT, department TEXT, topic TEXT DEFAULT 'Módulo 1 — Phishing', status TEXT DEFAULT 'Pendiente', score INTEGER DEFAULT 0, last_completed TEXT, UNIQUE(email, topic))""")
        c.execute("""CREATE TABLE IF NOT EXISTS remediation_tasks (id INTEGER PRIMARY KEY AUTOINCREMENT, organization_id INTEGER, scan_id INTEGER, hostname TEXT, finding_vector TEXT, severity TEXT DEFAULT 'MEDIO', status TEXT DEFAULT 'Pendiente', notes TEXT)""")
        try:
            c.execute("ALTER TABLE remediation_tasks ADD COLUMN organization_id INTEGER;")
            c.execute("ALTER TABLE remediation_tasks ADD COLUMN scan_id INTEGER;")
            c.execute("ALTER TABLE remediation_tasks ADD COLUMN severity TEXT;")
        except Exception: pass
        c.execute("""CREATE TABLE IF NOT EXISTS remediation_logs (id INTEGER PRIMARY KEY AUTOINCREMENT, task_id INTEGER, timestamp TEXT, status TEXT, notes TEXT)""")
        conn.commit()
    c.close()
    conn.close()

init_db()

def save_scan_to_db(hostname, ip, risk_score, findings_count, report_type_val, organization_id=None, findings=None):
    try:
        conn = get_db_connection()
        conn.autocommit = True
        c = conn.cursor()
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        is_pg = "postgres" in st.secrets
        ph = "%s" if is_pg else "?"
        
        if is_pg:
            if organization_id is not None:
                c.execute(f"INSERT INTO history (timestamp, hostname, ip, risk_score, findings_count, report_type, organization_id) VALUES ({ph}, {ph}, {ph}, {ph}, {ph}, {ph}, {ph}) RETURNING id", (timestamp, hostname, ip, risk_score, findings_count, report_type_val, organization_id))
            else:
                c.execute(f"INSERT INTO history (timestamp, hostname, ip, risk_score, findings_count, report_type) VALUES ({ph}, {ph}, {ph}, {ph}, {ph}, {ph}) RETURNING id", (timestamp, hostname, ip, risk_score, findings_count, report_type_val))
            scan_id = c.fetchone()[0]
        else:
            if organization_id is not None:
                c.execute(f"INSERT INTO history (timestamp, hostname, ip, risk_score, findings_count, report_type, organization_id) VALUES ({ph}, {ph}, {ph}, {ph}, {ph}, {ph}, {ph})", (timestamp, hostname, ip, risk_score, findings_count, report_type_val, organization_id))
            else:
                c.execute(f"INSERT INTO history (timestamp, hostname, ip, risk_score, findings_count, report_type) VALUES ({ph}, {ph}, {ph}, {ph}, {ph}, {ph})", (timestamp, hostname, ip, risk_score, findings_count, report_type_val))
            scan_id = c.lastrowid
            
        if findings:
            for f in findings:
                vec = f['vector']
                sev = f.get('severity', 'MEDIO')
                if organization_id is not None:
                    c.execute(f"INSERT INTO remediation_tasks (organization_id, scan_id, hostname, finding_vector, severity, status) VALUES ({ph}, {ph}, {ph}, {ph}, {ph}, 'Pendiente')", (organization_id, scan_id, hostname, vec, sev))
                else:
                    c.execute(f"INSERT INTO remediation_tasks (organization_id, scan_id, hostname, finding_vector, severity, status) VALUES (NULL, {ph}, {ph}, {ph}, {ph}, 'Pendiente')", (scan_id, hostname, vec, sev))
        c.close()
        conn.close()
    except Exception as e:
        st.error(f"Error detallado al guardar el escaneo en la base de datos: {e}")

def get_scan_history(org_id=None):
    conn = get_db_connection()
    is_pg = "postgres" in st.secrets
    ph = "%s" if is_pg else "?"
    if org_id is not None:
        query = f'SELECT timestamp AS "Fecha y Hora", hostname AS "Dominio / Host", ip AS "IP", risk_score AS "Risk Score (/100)", findings_count AS "Vulnerabilidades", report_type AS "Plantilla" FROM history WHERE organization_id = {ph} ORDER BY id DESC'
        df = pd.read_sql_query(query, conn, params=(org_id,))
    else:
        query = f'SELECT timestamp AS "Fecha y Hora", hostname AS "Dominio / Host", ip AS "IP", risk_score AS "Risk Score (/100)", findings_count AS "Vulnerabilidades", report_type AS "Plantilla" FROM history WHERE organization_id IS NULL ORDER BY id DESC'
        df = pd.read_sql_query(query, conn)
    conn.close()
    return df

def get_employees_df():
    conn = get_db_connection()
    df = pd.read_sql_query('SELECT email AS "Correo Electrónico", department AS "Departamento", topic AS "Campaña / Tema", status AS "Estado", score AS "Calificación (%)", last_completed AS "Última Evaluación" FROM employees', conn)
    conn.close()
    return df


# ==========================================
# 3. MÓDULOS DE ESCANEO Y SEGURIDAD
# ==========================================
def get_geolocation(hostname):
    geo_data = {"ip": "N/A", "country": "Desconocido", "city": "Desconocido", "org": "Desconocido"}
    try:
        ip = socket.gethostbyname(hostname)
        geo_data["ip"] = ip
        url = f"http://ip-api.com/json/{ip}?fields=status,country,city,org,isp"
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            data = response.json()
            if data.get("status") == "success":
                geo_data["country"] = data.get("country", "Desconocido")
                geo_data["city"] = data.get("city", "Desconocido")
                geo_data["org"] = data.get("org", data.get("isp", "Desconocido"))
    except Exception: pass
    return geo_data

def check_ssl_certificate(hostname):
    ssl_info = {"valid": False, "issuer": "Desconocido", "expires_soon": False, "days_remaining": 0, "details": "No se pudo verificar el certificado SSL/TLS."}
    try:
        context = ssl.create_default_context()
        with socket.create_connection((hostname, 443), timeout=5) as sock:
            with context.wrap_socket(sock, server_hostname=hostname) as ssock:
                cert = ssock.getpeercert()
                not_after_str = cert.get('notAfter')
                if not_after_str:
                    expires_date = datetime.datetime.strptime(not_after_str, "%b %d %H:%M:%S %Y %Z")
                    days_left = (expires_date - datetime.datetime.utcnow()).days
                    ssl_info["days_remaining"] = days_left
                    ssl_info["valid"] = True
                    issuer_dict = dict(x[0] for x in cert.get('issuer', ((('commonName', ''),),)) )
                    ssl_info["issuer"] = issuer_dict.get('commonName', issuer_dict.get('organizationName', 'Desconocido'))
                    if days_left < 30:
                        ssl_info["expires_soon"] = True
                        ssl_info["details"] = f"Certificado válido pero expira pronto ({days_left} días restantes)."
                    else:
                        ssl_info["details"] = f"Certificado SSL válido. Expira en {days_left} días."
    except Exception as e: ssl_info["details"] = f"Error al verificar SSL: {str(e)}"
    return ssl_info

def check_email_security(hostname):
    email_sec = {"spf": False, "dmarc": False}
    try:
        res_spf = requests.get(f"https://cloudflare-dns.com/dns-query?name={hostname}&type=TXT", headers={"Accept": "application/dns-json"}, timeout=4)
        if res_spf.status_code == 200:
            for ans in res_spf.json().get("Answer", []):
                if "v=spf1" in ans.get("data", ""): email_sec["spf"] = True
        res_dmarc = requests.get(f"https://cloudflare-dns.com/dns-query?name=_dmarc.{hostname}&type=TXT", headers={"Accept": "application/dns-json"}, timeout=4)
        if res_dmarc.status_code == 200:
            for ans in res_dmarc.json().get("Answer", []):
                if "v=DMARC1" in ans.get("data", ""): email_sec["dmarc"] = True
    except Exception: pass
    return email_sec

def discover_subdomains(domain):
    subdomains = set()
    try:
        url = f"https://crt.sh/?q=%25.{domain}&output=json"
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            data = response.json()
            for entry in data:
                for sub in entry.get("name_value", "").split("\n"):
                    sub = sub.strip().lower()
                    if domain in sub and "*" not in sub and "@" not in sub:
                        subdomains.add(sub)
    except Exception: pass
    return sorted(list(subdomains))[:12]

def scan_ports(hostname):
    common_ports = {21: "FTP", 22: "SSH", 80: "HTTP", 443: "HTTPS", 3306: "MySQL", 8080: "HTTP-Proxy", 8443: "HTTPS-Panel"}
    open_ports = []
    for port, service in common_ports.items():
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(1.0)
            if s.connect_ex((hostname, port)) == 0:
                open_ports.append({"port": port, "service": service})
            s.close()
        except Exception: pass
    return open_ports

def scan_target(url):
    parsed_url = urlparse(url)
    hostname = parsed_url.hostname or url.replace("https://", "").replace("http://", "").split("/")[0]
    
    findings = []
    stats = {"Críticas": 0, "Medias": 0, "Bajas": 0, "Seguras": 0}
    
    open_ports = scan_ports(hostname)
    subdomains = discover_subdomains(hostname)
    geo = get_geolocation(hostname)
    email_sec = check_email_security(hostname)
    ssl_info = check_ssl_certificate(hostname)
    
    if not ssl_info["valid"]:
        stats["Críticas"] += 1
        findings.append({"vector": "Certificado SSL/TLS Inválido o Ausente", "severity": "CRÍTICO", "badge": "badge-critical", "exec_title": "Fallo Crítico en Cifrado HTTPS", "desc": ssl_info["details"], "impact": "Los navegadores bloquearán el acceso a la web.", "fix": "Instalar certificado SSL/TLS válido.", "compliance": "PCI-DSS 4.1 / ISO 27001 / SOC 2", "snippet": f"certbot --nginx -d {hostname}"})
    elif ssl_info["expires_soon"]:
        stats["Medias"] += 1
        findings.append({"vector": f"Certificado SSL/TLS próximo a expirar ({ssl_info['days_remaining']} días)", "severity": "MEDIO", "badge": "badge-medium", "exec_title": "Riesgo de Expiración Próxima", "desc": ssl_info["details"], "impact": "Los servicios web dejarán de operar al caducar.", "fix": "Renovar el certificado.", "compliance": "ISO 27001 A.12.1", "snippet": "certbot renew --dry-run"})
    else:
        stats["Seguras"] += 1

    if email_sec["spf"]: stats["Seguras"] += 1
    else:
        stats["Medias"] += 1
        findings.append({"vector": "Ausencia de Registro SPF (Phishing / GDPR)", "severity": "MEDIO", "badge": "badge-medium", "exec_title": "Vulnerabilidad en Postura de Correo", "desc": "El dominio carece de un registro SPF válido.", "impact": "Facilita la suplantación de identidad (phishing).", "fix": "Publicar registro TXT con directivas SPF.", "compliance": "ISO 27001 A.13.2 / GDPR", "snippet": f'{hostname}. 3600 IN TXT "v=spf1 include:_spf.example.com ~all"'})

    if email_sec["dmarc"]: stats["Seguras"] += 1
    else:
        stats["Medias"] += 1
        findings.append({"vector": "Ausencia de Política DMARC", "severity": "MEDIO", "badge": "badge-medium", "exec_title": "Falta de Control DMARC", "desc": "El dominio carece de una política DMARC.", "impact": "Pérdida de visibilidad sobre intentos de fraude.", "fix": "Configurar registro TXT en _dmarc.", "compliance": "ISO 27001 A.13.1", "snippet": f'_dmarc.{hostname}. 3600 IN TXT "v=DMARC1; p=reject;"'})

    for p in open_ports:
        if p['port'] in [21, 3306, 8080, 8443]:
            stats["Medias"] += 1
            findings.append({"port": p['port'], "service": p['service'], "vector": f"Puerto {p['port']} ({p['service']}) Abierto al Público", "severity": "MEDIO", "badge": "badge-medium", "exec_title": f"Servicio Expuesto en Puerto {p['port']}", "desc": f"El puerto {p['port']} es accesible desde internet.", "impact": "Expuesto a ataques de fuerza bruta.", "fix": "Restringir el acceso mediante Firewall.", "compliance": "PCI-DSS 1.3 / SOC 2 CC6.1", "snippet": f"sudo ufw deny {p['port']}/tcp"})

    try:
        response = requests.get(url, timeout=10)
        headers = response.headers
        if "Strict-Transport-Security" in headers: stats["Seguras"] += 1
        else:
            stats["Críticas"] += 1
            findings.append({"vector": "HTTP Strict Transport Security (HSTS) Ausente", "severity": "CRÍTICO", "badge": "badge-critical", "exec_title": "Ausencia de HSTS", "desc": "La cabecera HSTS no está configurada.", "impact": "Riesgo de intercepción de tráfico.", "fix": "Configurar la cabecera HSTS.", "compliance": "PCI-DSS 4.1 / HIPAA", "snippet": 'add_header Strict-Transport-Security "max-age=31536000;" always;'})
        
        if "Content-Security-Policy" in headers: stats["Seguras"] += 1
        else:
            stats["Medias"] += 1
            findings.append({"vector": "Content Security Policy (CSP) Ausente", "severity": "MEDIO", "badge": "badge-medium", "exec_title": "Ausencia de CSP", "desc": "No se detectó la cabecera Content-Security-Policy.", "impact": "Riesgo de ataques XSS (Cross-Site Scripting).", "fix": "Implementar directivas CSP robustas.", "compliance": "OWASP / SOC 2", "snippet": 'add_header Content-Security-Policy "default-src \'self\';";'})
        
        if "X-Content-Type-Options" in headers: stats["Seguras"] += 1
        else:
            stats["Bajas"] += 1
            findings.append({"vector": "Cabecera X-Content-Type-Options Ausente", "severity": "BAJO", "badge": "badge-low", "exec_title": "MIME-Sniffing Risk", "desc": "Falta la protección contra sniffing de tipos MIME.", "impact": "Interpretación incorrecta de archivos por el navegador.", "fix": "Añadir X-Content-Type-Options nosniff.", "compliance": "OWASP Top 10", "snippet": 'add_header X-Content-Type-Options "nosniff";'})
    except Exception: pass

    penalty = (stats["Críticas"] * 25) + (stats["Medias"] * 10) + (stats["Bajas"] * 5)
    risk_score = max(0, 100 - penalty)
    return findings, stats, open_ports, hostname, subdomains, geo, email_sec, ssl_info, risk_score


# ==========================================
# 4. GENERACIÓN DE REPORTES Y ALERTAS
# ==========================================
def send_webhook_alert(webhook_url, hostname, risk_score, findings_count):
    if not webhook_url: return
    try:
        payload = {"text": f"🚨 *CyberAudits Security Alert*\n• Objetivo: {hostname}\n• Risk Score: *{risk_score}/100*\n• Vulnerabilidades detectadas: *{findings_count}*"}
        requests.post(webhook_url, json=payload, timeout=5)
    except Exception: pass

def generate_chart(stats):
    labels, sizes, colors = list(stats.keys()), list(stats.values()), ['#dc2626', '#f59e0b', '#3b82f6', '#10b981']
    non_zero_data = [(l, s, c) for l, s, c in zip(labels, sizes, colors) if s > 0]
    if not non_zero_data: non_zero_data = [("Seguras", 1, "#10b981")]
    l_filt, s_filt, c_filt = zip(*non_zero_data)
    fig, ax = plt.subplots(figsize=(4.5, 2.8))
    ax.pie(s_filt, labels=l_filt, colors=c_filt, autopct='%1.1f%%', startangle=90, textprops={'fontsize': 8.5, 'weight': 'bold'})
    ax.axis('equal')
    plt.title("Distribución de Riesgos en la Infraestructura", fontsize=9.5, fontweight='bold', color="#1e293b")
    plt.tight_layout()
    chart_path = "vulnerability_chart.png"
    plt.savefig(chart_path, dpi=300, bbox_inches='tight', transparent=True)
    plt.close()
    with open(chart_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")

def generate_docx(hostname, geo, email_sec, ssl_info, open_ports, subdomains, findings, risk_score, agency_name, agency_tagline, report_type, recipient_name, report_subject):
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
        
    run_title = doc.add_paragraph().add_run(f"INFORME: {report_type.upper()}")
    run_title.font.size, run_title.font.bold, run_title.font.color.rgb = Pt(15), True, RGBColor(15, 23, 42)
    
    run_sub = doc.add_paragraph().add_run(f"Emitido por: {agency_name} ({agency_tagline})\nObjetivo analizado: {hostname} | Risk Score: {risk_score}/100")
    run_sub.font.size, run_sub.font.color.rgb = Pt(10), RGBColor(100, 116, 139)
    
    doc.add_heading("1. Datos Generales y Metadatos del Objetivo", level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ("Dominio / Hostname", hostname), ("Dirección IP", geo['ip']),
        ("Risk Score Global", f"{risk_score} / 100"), ("Ubicación Geográfica", f"{geo['city']}, {geo['country']} ({geo['org']})"),
        ("Seguridad de Correo", f"SPF: {'OK' if email_sec['spf'] else 'Ausente'} | DMARC: {'OK' if email_sec['dmarc'] else 'Ausente'}"),
        ("Certificado SSL/TLS", f"{ssl_info['details']}")
    ]
    for i, (k, v) in enumerate(data): table.cell(i, 0).text, table.cell(i, 1).text = k, str(v)
        
    doc.add_heading("2. Detalle de Hallazgos y Guía de Remediación", level=2)
    for idx, f in enumerate(findings, 1):
        run_h = doc.add_paragraph().add_run(f"#{idx} - {f['vector']} [{f['severity']}] | Norma: {f.get('compliance', 'N/A')}")
        run_h.font.bold, run_h.font.size = True, Pt(11)
        doc.add_paragraph(f"Descripción: {f['desc']}")
        doc.add_paragraph(f"Impacto de Negocio: {f['impact']}")
        p_fix = doc.add_paragraph()
        p_fix.add_run("Remediación: ").font.bold = True
        p_fix.add_run(f"{f['fix']}")
        if "snippet" in f:
            p_snip = doc.add_paragraph()
            p_snip.add_run("Configuración sugerida:\n").font.bold = True
            run_code = p_snip.add_run(f"{f['snippet']}")
            run_code.font.name, run_code.font.size = 'Courier New', Pt(9.5)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_pdf(url, findings, stats, chart_base64, open_ports, hostname, subdomains, geo, email_sec, ssl_info, risk_score, agency_name, agency_tagline, report_type, recipient_name, report_subject, logo_b64, output_filename):
    logo_html = f'<img src="data:image/png;base64,{logo_b64}" style="max-height: 55px; width: auto; float: right; margin-top: 2px;" alt="Logo">' if logo_b64 else ""
    ports_html = "".join([f"<tr><td><code>{p['port']}</code></td><td><strong>{p['service']}</strong> (Open)</td></tr>" for p in open_ports]) or "<tr><td colspan='2' style='text-align:center;'>No ports detected.</td></tr>"
    sub_html = "".join([f"<li><code>{sub}</code></li>" for sub in subdomains]) or "<li>No subdomains found.</li>"
    spf_badge = "<span style='color:green;'><b>OK</b></span>" if email_sec["spf"] else "<span style='color:red;'><b>Ausente</b></span>"
    dmarc_badge = "<span style='color:green;'><b>OK</b></span>" if email_sec["dmarc"] else "<span style='color:red;'><b>Ausente</b></span>"
    ssl_badge = "<span style='color:green;'><b>Válido</b></span>" if ssl_info["valid"] and not ssl_info["expires_soon"] else "<span style='color:orange;'><b>Revisar</b></span>"
    
    items_html_full = ""
    for idx, f in enumerate(findings, 1):
        snippet_box = f"<pre style=\"background:#f1f5f9;padding:6px;border-radius:4px;font-size:7pt;color:#0369a1;overflow-x:auto;\"><code>{f.get('snippet', '')}</code></pre>" if "snippet" in f else ""
        items_html_full += f"""
        <div class="finding-card">
            <div class="finding-header"><span class="finding-num">#{idx}</span><span class="finding-title">{f['vector']}</span><span class="{f['badge']}">{f['severity']}</span></div>
            <div class="finding-body"><p><strong>Norma:</strong> <code>{f.get('compliance', 'N/A')}</code></p><p><strong>Descripción:</strong> {f['desc']}</p><p><strong>Impacto:</strong> {f['impact']}</p><div class="solution-box"><p><strong>Remediación:</strong> <code>{f['fix']}</code></p></div>{snippet_box}</div>
        </div>
        """
        
    content_html = f"""
    <div class="header-banner">
        <div class="banner-left">
            <h1>{report_type}</h1>
            <p>Elaborado por: <strong>{agency_name}</strong> ({agency_tagline})</p>
        </div>
        <div class="banner-right">{logo_html}</div>
    </div>
    <table style="width: 100%; margin-bottom: 6px; border: none;">
        <tr>
            <td style="border: none; width: 25%;"><div class="meta-item"><div class="meta-label">Objetivo</div><div class="meta-value">{hostname}</div></div></td>
            <td style="border: none; width: 25%;"><div class="meta-item"><div class="meta-label">Risk Score</div><div class="meta-value" style="color: {'#10b981' if risk_score > 70 else '#f59e0b' if risk_score > 40 else '#dc2626'};">{risk_score} / 100</div></div></td>
            <td style="border: none; width: 25%;"><div class="meta-item"><div class="meta-label">SSL</div><div class="meta-value">{ssl_badge}</div></div></td>
            <td style="border: none; width: 25%;"><div class="meta-item"><div class="meta-label">SPF / DMARC</div><div class="meta-value">{spf_badge} / {dmarc_badge}</div></div></td>
        </tr>
    </table>
    <h2>1. Visión e Infraestructura</h2>
    <div class="executive-box"><p style="margin:0;">Risk Score corporativo: <strong>{risk_score}/100</strong>.</p></div>
    <table style="width: 100%; border: none; margin-bottom: 6px;">
        <tr>
            <td style="width: 50%; vertical-align: top; border: none;">
                <div class="card"><h3 style="margin:0; font-size:8.5pt;">Puertos Críticos:</h3><table class="ports-table"><thead><tr><th>Puerto</th><th>Servicio</th></tr></thead><tbody>{ports_html}</tbody></table></div>
            </td>
            <td style="width: 50%; vertical-align: top; border: none;">
                <div class="card"><h3 style="margin:0; font-size:8.5pt;">Subdominios:</h3><ul style="margin:0; padding-left:14px; font-size:7pt; max-height:75px; overflow:hidden;">{sub_html}</ul></div>
            </td>
        </tr>
    </table>
    <div class="card" style="text-align: center; padding: 4px;"><div class="chart-container"><img src="data:image/png;base64,{chart_base64}" alt="Gráfico"></div></div>
    <div style="page-break-after: always;"></div>
    <h2>2. Hallazgos Detallados</h2>
    {items_html_full}
    """

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{ size: A4; margin: 10mm 12mm; background-color: #f8fafc; @bottom-right {{ content: "Page " counter(page) " of " counter(pages); font-size: 8pt; color: #64748b; }} }}
            body {{ font-family: Helvetica, Arial, sans-serif; color: #334155; margin: 0; padding: 0; font-size: 8.5pt; line-height: 1.35; }}
            .header-banner {{ background: #0f172a; color: white; padding: 10px 15px; border-radius: 5px; margin-bottom: 6px; overflow: hidden; }}
            .banner-left {{ float: left; width: 70%; }}
            .banner-right {{ float: right; width: 28%; text-align: right; }}
            .header-banner h1 {{ margin: 0; font-size: 13pt; }}
            .header-banner p {{ margin: 0; color: #94a3b8; font-size: 8pt; }}
            .meta-item {{ background: white; padding: 4px 8px; border: 1px solid #e2e8f0; border-radius: 4px; }}
            .meta-label {{ font-size: 6pt; color: #64748b; text-transform: uppercase; }}
            .meta-value {{ font-size: 8pt; font-weight: 600; color: #0f172a; }}
            h2 {{ color: #0f172a; font-size: 9.5pt; border-left: 3px solid #3b82f6; padding-left: 5px; margin-top: 6px; margin-bottom: 4px; }}
            .card {{ background: white; border: 1px solid #e2e8f0; border-radius: 5px; padding: 6px 8px; margin-bottom: 5px; }}
            .badge-critical {{ background-color: #fee2e2; color: #991b1b; padding: 2px 4px; border-radius: 3px; font-size: 6.5pt; float: right; }}
            .badge-medium {{ background-color: #fef3c7; color: #92400e; padding: 2px 4px; border-radius: 3px; font-size: 6.5pt; float: right; }}
            .badge-low {{ background-color: #dbeafe; color: #1e40af; padding: 2px 4px; border-radius: 3px; font-size: 6.5pt; float: right; }}
            .chart-container {{ text-align: center; }}
            .chart-container img {{ max-width: 65%; height: auto; }}
            .executive-box {{ background-color: #eff6ff; border-left: 3px solid #3b82f6; padding: 5px 8px; margin-bottom: 5px; }}
            table.ports-table {{ width: 100%; border-collapse: collapse; font-size: 7.5pt; }}
            table.ports-table th {{ background-color: #f1f5f9; padding: 2px; border-bottom: 2px solid #cbd5e1; text-align: left; }}
            table.ports-table td {{ padding: 2px; border-bottom: 1px solid #e2e8f0; }}
            .finding-card {{ background: white; border: 1px solid #cbd5e1; border-radius: 4px; margin-bottom: 5px; page-break-inside: avoid; }}
            .finding-header {{ background-color: #f1f5f9; padding: 4px 6px; border-bottom: 1px solid #cbd5e1; overflow: hidden; }}
            .finding-title {{ font-weight: bold; color: #0f172a; font-size: 8pt; }}
            .finding-body {{ padding: 5px 6px; }}
            .solution-box {{ background-color: #f8fafc; border-left: 3px solid #0284c7; padding: 4px 6px; margin-top: 3px; }}
            .solution-box code {{ color: #0369a1; font-size: 7pt; }}
            .disclaimer {{ font-size: 7pt; color: #64748b; margin-top: 6px; text-align: center; font-style: italic; }}
        </style>
    </head>
    <body>{content_html}<div class="disclaimer">Nota: Evaluaciones perimetrales externas en tiempo real.</div></body>
    </html>
    """
    HTML(string=html_content).write_pdf(output_filename)


# ==========================================
# 5. CONSTANTES Y DATOS DE CONCIENCIACIÓN
# ==========================================
TRAINING_TOPICS = {
    "Módulo 1 — Phishing": {
        "title": "Módulo 1 — Phishing y Detección de Fraudes",
        "theory": "### 🎣 ¿Qué es el Phishing?\nEl **phishing** es una técnica utilizada para engañar a las personas y obtener información confidencial...\n#### 1. Señales de Alerta y Urgencia\n* **Situaciones sospechosas:** Un mensaje que nos pide actuar inmediatamente es una señal clara.\n* **Uso de la urgencia:** Los ciberdelincuentes suelen generar urgencia.\n#### 2. Cómo Actuar\n* **Correos bancarios:** Nunca hagas clic inmediatamente.\n* **Información crítica:** Nunca proporcionar contraseñas.\n* **Archivos adjuntos:** Comprobar si el mensaje es legítimo.\n* **Errores:** Cambiar la contraseña inmediatamente desde la página oficial.\n* **Ámbito laboral:** Informarlo siguiendo el procedimiento de la organización.",
        "questions": [
            {"q": "1. ¿Qué es el phishing?", "options": ["Programa protector.", "Técnica para engañar y obtener información.", "Sistema de Internet."], "correct": 1},
            {"q": "2. ¿Cuál es una señal de phishing?", "options": ["Mensaje urgente amenazando bloqueo.", "Mensaje esperado.", "Notificación habitual."], "correct": 0},
            {"q": "3. Recibes un correo de tu banco con un enlace. ¿Qué haces?", "options": ["Clic inmediato.", "Comprobar en la app oficial.", "Responder solicitando info."], "correct": 1},
            {"q": "4. ¿Qué nunca debemos dar por enlace?", "options": ["Contraseñas o códigos.", "Ciudad.", "Idioma."], "correct": 0},
            {"q": "5. ¿Por qué generan urgencia?", "options": ["Más tiempo para analizar.", "Para que la persona actúe rápido sin comprobar.", "Mejorar seguridad."], "correct": 1},
            {"q": "6. Archivo adjunto inesperado. ¿Qué hacer?", "options": ["Abrirlo.", "Reenviarlo.", "Comprobar remitente."], "correct": 2},
            {"q": "7. Introdujiste tu contraseña en página falsa. ¿Qué hacer?", "options": ["Nada.", "Cambiar contraseña desde sitio oficial e informar.", "Pedir ayuda a compañero."], "correct": 1},
            {"q": "8. Correo sospechoso en trabajo. ¿Qué hacer?", "options": ["Reenviar a todos.", "Clic al enlace.", "Informar por procedimiento de la empresa."], "correct": 2}
        ]
    },
    "Módulo 2 — Contraseñas seguras": {
        "title": "Módulo 2 — Contraseñas Seguras y MFA",
        "theory": "### 🔑 Gestión de Contraseñas y Autenticación\nUna contraseña segura es larga y difícil de adivinar...\n#### 1. Buenas Prácticas\n* **Info personal:** Evitar usar nombres/fechas.\n* **Reutilización:** No usar la misma para todo.\n* **Dónde guardarlas:** Usar gestor de contraseñas.\n#### 2. Autenticación Multifactor (MFA)\n* **¿Qué es?:** Medida adicional además de la contraseña.\n* **Códigos inesperados:** No compartirlos con nadie.",
        "questions": [
            {"q": "1. Característica de contraseña segura", "options": ["Larga y difícil.", "Nombre + 123.", "Misma en todas partes."], "correct": 0},
            {"q": "2. Contraseña menos segura", "options": ["Frase larga.", "Palabras/caracteres.", "123456."], "correct": 2},
            {"q": "3. ¿Por qué no reutilizar contraseñas?", "options": ["Si una se expone, otras quedan en riesgo.", "Funcionan una vez.", "Internet lento."], "correct": 0},
            {"q": "4. A evitar al crear contraseña", "options": ["Info personal fácil.", "Contraseña larga.", "Diferente por cuenta."], "correct": 0},
            {"q": "5. ¿Qué es MFA?", "options": ["Elimina contraseñas.", "Comprobación adicional.", "Aumenta velocidad."], "correct": 1},
            {"q": "6. Código SMS no solicitado. ¿Qué hacer?", "options": ["Compartirlo.", "Publicarlo.", "No compartir y revisar actividad."], "correct": 2},
            {"q": "7. ¿Dónde evitar guardar contraseñas?", "options": ["Gestor confiable.", "Papel en el monitor.", "Sistema de organización."], "correct": 1},
            {"q": "8. ¿Qué hacer si hay MFA disponible?", "options": ["Activarlo.", "Desactivarlo.", "Compartirlo."], "correct": 0}
        ]
    },
    "Módulo 3 — Seguridad en el puesto de trabajo": {
        "title": "Módulo 3 — Seguridad Física",
        "theory": "### 🏢 Seguridad en el Entorno Laboral\nLa protección no solo depende del software...\n#### 1. Bloqueo\n* **Bloqueo pantalla:** Windows + L al alejarse.\n* **USB desconocidos:** Entregarlos, no conectarlos.\n#### 2. Información\n* **Documentos:** Protegerlos de accesos no autorizados.\n* **Archivos compañeros:** Comprobar antes de abrir.\n* **Fuera de oficina:** Cuidar dispositivos por robo/pérdida.",
        "questions": [
            {"q": "1. Al alejarse del ordenador", "options": ["Dejar abierto.", "Bloquear pantalla.", "Escribir contraseña."], "correct": 1},
            {"q": "2. Teclas bloqueo Windows", "options": ["Windows + L.", "Ctrl + C.", "Alt + F4."], "correct": 0},
            {"q": "3. USB desconocido en empresa", "options": ["Conectar.", "Llevar a casa.", "Entregar sin conectar."], "correct": 2},
            {"q": "4. Peligro de USB", "options": ["Malware.", "Siempre vacíos.", "Velocidad."], "correct": 0},
            {"q": "5. Documentos confidenciales", "options": ["En escritorio.", "Proteger acceso.", "Fotografiar."], "correct": 1},
            {"q": "6. Archivo inesperado de compañero", "options": ["Abrir.", "Comprobar envío.", "Reenviar."], "correct": 1},
            {"q": "7. Buena práctica", "options": ["Instalar todo.", "Compartir sesión.", "Actualizar y seguir políticas."], "correct": 2},
            {"q": "8. Cuidado fuera de oficina", "options": ["Riesgo pérdida/robo de info.", "Lentitud.", "No funcionan."], "correct": 0}
        ]
    },
    "Módulo 4 — Vishing y Smishing": {
        "title": "Módulo 4 — Llamadas/Mensajes Falsos",
        "theory": "### ☎️ Fraudes Móviles\n#### 1. Conceptos\n* **Vishing:** Estafa telefónica.\n* **Smishing:** Estafa por SMS.\n#### 2. Detección\n* **Llamadas/códigos:** Nunca compartir códigos de SMS.\n* **SMS de paquetes:** No clicar enlaces.\n* **Familiares:** Comprobar identidad si piden dinero.\n* **Acción:** Finalizar llamada sospechosa.",
        "questions": [
            {"q": "1. ¿Qué es vishing?", "options": ["Estafa telefónica.", "Antivirus.", "Proteger contraseñas."], "correct": 0},
            {"q": "2. ¿Qué es smishing?", "options": ["Backup.", "Estafa SMS.", "Mejorar Wi-Fi."], "correct": 1},
            {"q": "3. Llamada banco pidiendo código SMS", "options": ["Dar código.", "Dar números.", "No compartir y verificar canal oficial."], "correct": 2},
            {"q": "4. Señal fraude telefónico", "options": ["Presión para actuar rápido.", "Permitir comprobación.", "No pide info."], "correct": 0},
            {"q": "5. SMS paquete con enlace", "options": ["Clic inmediato.", "Comprobar en página oficial de empresa.", "Dar tarjeta."], "correct": 1},
            {"q": "6. Familiar desconocido pide dinero", "options": ["Enviar dinero.", "Pedir contraseña.", "Comprobar identidad por otro medio."], "correct": 2},
            {"q": "7. A evitar en llamada", "options": ["Contraseñas/datos bancarios.", "Nombre de empresa público.", "La hora."], "correct": 0},
            {"q": "8. Acción ante llamada sospechosa", "options": ["Hablar.", "Finalizar y contactar número oficial.", "Dar info falsa."], "correct": 1}
        ]
    }
}


# ==========================================
# 6. INTERFAZ PRINCIPAL DE LA APLICACIÓN (UI)
# ==========================================
query_params = st.query_params
employee_token = query_params.get("empleado")
topic_token = query_params.get("tema", "Módulo 1 — Phishing")

if employee_token:
    # --- PORTAL DE COLABORADORES (CONCIENCIACIÓN) ---
    st.markdown("""
        <div class="employee-portal-banner">
            <h2>🎓 Portal Corporativo de Concienciación en Ciberseguridad</h2>
            <p>Capacitación esencial para colaboradores</p>
        </div>
        """, unsafe_allow_html=True)
    
    selected_topic_data = TRAINING_TOPICS.get(topic_token, TRAINING_TOPICS["Módulo 1 — Phishing"])
    st.info(f"👤 Colaborador: **{employee_token}** | Módulo Asignado: **{topic_token}**")
    
    conn = get_db_connection()
    c = conn.cursor()
    placeholder = "%s" if "postgres" in st.secrets else "?"
    c.execute(f"SELECT status, score FROM employees WHERE email = {placeholder} AND topic = {placeholder}", (employee_token, topic_token))
    row = c.fetchone()
    
    if not row:
        try:
            c.execute(f"INSERT INTO employees (email, department, topic, status) VALUES ({placeholder}, {placeholder}, {placeholder}, {placeholder})", (employee_token, "General", topic_token, "Pendiente"))
            conn.commit()
        except Exception: pass
        current_status, current_score = "Pendiente", 0
    else:
        current_status, current_score = row
    c.close()
    conn.close()
    
    if current_status == "Completado":
        st.success(f"✅ ¡Ya hay un registro de examen completado para **{topic_token}** con una calificación de **{current_score}%**!")
    else:
        st.markdown(f"### 📚 Material de Estudio: {topic_token}")
        st.markdown(selected_topic_data["theory"])
        st.markdown("---")
        st.markdown("### 📝 Cuestionario de Evaluación")
        
        with st.form("employee_deep_quiz_form"):
            user_answers = {}
            for idx, q_item in enumerate(selected_topic_data["questions"]):
                st.write(f"**{q_item['q']}**")
                user_choice = st.radio("Seleccione una opción:", q_item["options"], key=f"q_{idx}", label_visibility="collapsed")
                user_answers[idx] = (user_choice, q_item["correct"])
                st.markdown("")
                
            if st.form_submit_button("Enviar Examen y Registrar Resultados"):
                score_points = sum(1 for idx, (chosen_text, correct_idx) in user_answers.items() if chosen_text == selected_topic_data["questions"][idx]["options"][correct_idx])
                final_score = int((score_points / len(selected_topic_data["questions"])) * 100)
                timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                
                conn = get_db_connection()
                c = conn.cursor()
                c.execute(f"UPDATE employees SET status = 'Completado', score = {placeholder}, last_completed = {placeholder} WHERE email = {placeholder} AND topic = {placeholder}", (final_score, timestamp, employee_token, topic_token))
                conn.commit()
                c.close()
                conn.close()
                
                st.success(f"🎉 ¡Examen enviado con éxito! Has obtenido una calificación de **{final_score}%** ({score_points}/{len(selected_topic_data['questions'])} aciertos).")
                st.rerun()

else:
    # --- PANEL PRINCIPAL ENTERPRISE (ADMIN) ---
    if "scanned" not in st.session_state: st.session_state.scanned = False
    if "failed_attempts" not in st.session_state: st.session_state.failed_attempts = 0
    if "org_success_msg" not in st.session_state: st.session_state.org_success_msg = ""
        
    st.markdown('<div class="enterprise-banner">🚀 <strong>CyberAudits Enterprise Suite:</strong> Plataforma perimetral de consultoría activa.</div>', unsafe_allow_html=True)
    st.title("🛡️ CyberAudits - Suite Enterprise")
    st.write("Plataforma integral de ciberseguridad: Auditoría perimetral y gestión de campañas de concienciación.")
    
    st.sidebar.header("🧭 Módulos de la Plataforma")
    modules_list = ["Auditoría Perimetral"]
    is_admin = False
    
    if st.session_state.failed_attempts >= 5:
        st.sidebar.error("⚠️ Acceso bloqueado temporalmente.")
    else:
        admin_password_input = st.sidebar.text_input("🔑 Contraseña Administrador", type="password")
        if admin_password_input:
            if hashlib.sha256(admin_password_input.encode()).hexdigest() == "b1db078a7a989c545804a3ed56cc961d11c35885cb3848dffaff39a2ea6b468e":
                is_admin = True
                st.session_state.failed_attempts = 0
            else:
                st.session_state.failed_attempts += 1
                st.sidebar.error("Contraseña incorrecta.")
                
    if is_admin: modules_list.append("🎓 Concienciación (Privado)")
        
    selected_module = st.sidebar.radio("Módulo Disponible", modules_list)
    st.sidebar.markdown("---")
    
    selected_org_id = None
    selected_org_name = "General / Sin Asignar"
    
    if selected_module == "Auditoría Perimetral":
        st.sidebar.header("🏢 Organización / Cliente")
        try:
            conn_org = get_db_connection()
            org_df = pd.read_sql_query("SELECT id, name FROM organizations", conn_org)
            conn_org.close()
        except Exception: org_df = pd.DataFrame(columns=["id", "name"])
            
        org_options = {"General / Sin Asignar": None}
        for _, row in org_df.iterrows(): org_options[row["name"]] = row["id"]
                
        selected_org_name = st.sidebar.selectbox("Cliente Objetivo", list(org_options.keys()))
        selected_org_id = org_options[selected_org_name]
        
        with st.sidebar.expander("➕ Añadir Nueva Organización"):
            with st.form("add_org_form", clear_on_submit=True):
                new_org_input = st.text_input("Nombre del Cliente")
                if st.form_submit_button("Guardar Cliente") and new_org_input:
                    try:
                        conn_add = get_db_connection()
                        c_add = conn_add.cursor()
                        c_add.execute(f"INSERT INTO organizations (name) VALUES ({'%s' if 'postgres' in st.secrets else '?'})", (new_org_input,))
                        conn_add.commit()
                        c_add.close()
                        conn_add.close()
                        st.session_state.org_success_msg = f"Cliente registrado: {new_org_input}"
                        st.rerun()
                    except Exception: st.warning("Organización existente o error.")
                        
        if st.session_state.org_success_msg:
            st.sidebar.success(st.session_state.org_success_msg)
            st.session_state.org_success_msg = ""
            
        if selected_org_id is not None and st.sidebar.button("🗑️ Eliminar Cliente"):
            try:
                conn_del = get_db_connection()
                c_del = conn_del.cursor()
                c_del.execute(f"DELETE FROM organizations WHERE id = {'%s' if 'postgres' in st.secrets else '?'}", (selected_org_id,))
                conn_del.commit()
                c_del.close()
                conn_del.close()
                st.sidebar.success(f"Cliente '{selected_org_name}' eliminado.")
                st.rerun()
            except Exception as e: st.sidebar.error(f"Error: {e}")
                    
        st.sidebar.markdown("---")
        st.sidebar.header("⚙️ Configuración de Informe")
        agency_name = st.sidebar.text_input("Agencia", value="SecOps Global Partners")
        agency_tagline = st.sidebar.text_input("Subtítulo", value="Consultoría y Ciberseguridad")
        logo_file = st.sidebar.file_uploader("Logo", type=["png", "jpg", "jpeg"])
        report_type = st.sidebar.selectbox("Plantilla", ["Informe Técnico Exhaustivo", "Informe Narrativo", "Normativa (ISO/Compliance)"])
        recipient_name = st.sidebar.text_input("Dirigido a", value="Dirección General")
        report_subject = st.sidebar.text_input("Asunto", value="Evaluación de Riesgos")
        st.sidebar.markdown("---")
        webhook_url_input = st.sidebar.text_input("Webhook URL (Slack/Teams)", type="password")
    else:
        agency_name, agency_tagline, logo_file, report_type, recipient_name, report_subject, selected_org_id, webhook_url_input = "SecOps Global", "Consultoría", None, "Informe Técnico", "Dirección", "Riesgos", None, ""
        
    st.sidebar.markdown("---")
    st.sidebar.caption("CyberAudits Enterprise v6.5 • Cloud Platform")
    
    if is_admin and selected_module == "🎓 Concienciación (Privado)":
        st.markdown("## 🎓 Gestión de Campañas y Directorio")
        sub_tab1, sub_tab2, sub_tab3 = st.tabs(["👥 Registro y CSV", "📊 Dashboard", "🔗 Enlaces"])
        with sub_tab1:
            st.markdown("Funciones de importación y gestión AD en desarrollo...")
        with sub_tab2:
            st.dataframe(get_employees_df(), use_container_width=True)
        with sub_tab3:
            st.info("Exportador de enlaces en construcción.")
    else:
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["🔍 Perimeter Scan", "📊 Security Analytics", "📜 Historial", "🛠️ Ticketera", "ℹ️ About"])
        
        with tab1:
            st.markdown("### 🎯 Quick Test Targets")
            col_btn1, col_btn2, col_btn3 = st.columns(3)
            quick_domain = "example.com" if col_btn1.button("🌐 example.com") else "scanme.nmap.org" if col_btn2.button("🌐 scanme.nmap.org") else "testphp.vulnweb.com" if col_btn3.button("🌐 testphp.vulnweb.com") else ""
            target_url = st.text_input("URL Objetivo", value=quick_domain or "https://")
            
            if st.button("🚀 Ejecutar Análisis"):
                if not target_url or target_url == "https://": st.error("URL inválida.")
                else:
                    if not target_url.startswith("http"): target_url = "https://" + target_url
                    with st.status("🔍 Analizando infraestructura...", expanded=True) as status:
                        findings, stats, open_ports, hostname, subdomains, geo, email_sec, ssl_info, risk_score = scan_target(target_url)
                        save_scan_to_db(hostname, geo["ip"], risk_score, len(findings), report_type, selected_org_id, findings)
                        if webhook_url_input: send_webhook_alert(webhook_url_input, hostname, risk_score, len(findings))
                        
                        chart_b64 = generate_chart(stats)
                        logo_b64 = base64.b64encode(logo_file.getvalue()).decode("utf-8") if logo_file else ""
                        pdf_filename = f"auditoria_{hostname}.pdf"
                        generate_pdf(target_url, findings, stats, chart_b64, open_ports, hostname, subdomains, geo, email_sec, ssl_info, risk_score, agency_name, agency_tagline, report_type, recipient_name, report_subject, logo_b64, pdf_filename)
                        docx_bytes = generate_docx(hostname, geo, email_sec, ssl_info, open_ports, subdomains, findings, risk_score, agency_name, agency_tagline, report_type, recipient_name, report_subject)
                        
                        status.update(label="✅ ¡Análisis completado!", state="complete", expanded=False)
                        st.session_state.update(scanned=True, findings=findings, hostname=hostname, geo=geo, email_sec=email_sec, ssl_info=ssl_info, risk_score=risk_score, pdf_filename=pdf_filename, docx_bytes=docx_bytes)

            if st.session_state.scanned:
                g1, g2, g3 = st.columns(3)
                g1.metric("Risk Score", f"{st.session_state.risk_score} / 100")
                g2.metric("IP", st.session_state.geo["ip"])
                g3.metric("SSL", "Válido" if st.session_state.ssl_info["valid"] else "Revisar")
                
                col_dl1, col_dl2 = st.columns(2)
                with col_dl1:
                    if os.path.exists(st.session_state.pdf_filename):
                        with open(st.session_state.pdf_filename, "rb") as pdf_file:
                            st.download_button("📥 PDF Ejecutivo", pdf_file, file_name=st.session_state.pdf_filename, mime="application/pdf", type="primary")
                with col_dl2:
                    if "docx_bytes" in st.session_state:
                        st.download_button("📝 DOCX Editable", st.session_state.docx_bytes, file_name=f"auditoria_{st.session_state.hostname}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")

        with tab2:
            if st.session_state.scanned:
                for f in st.session_state.findings:
                    with st.expander(f"📌 {f['vector']} [{f['severity']}]"):
                        st.write(f"**Descripción:** {f['desc']}")
                        st.write(f"**Remediación:** {f['fix']}")
            else: st.info("Ejecuta un escaneo primero.")

        with tab3:
            history_df = get_scan_history(org_id=selected_org_id)
            if not history_df.empty: st.dataframe(history_df, use_container_width=True)
            else: st.info("No hay historial para este cliente.")

        with tab4:
            st.write("Selecciona un escaneo para gestionar la remediación (Ticketera).")
            # --- Lógica de la ticketera omitida para brevedad de visualización ---
            st.info("Dirígete al análisis principal para habilitar la generación de tickets.")

        with tab5:
            st.markdown("**CyberAudits Enterprise Suite** es una plataforma integral orientada a consultorías de ciberseguridad corporativa.")
